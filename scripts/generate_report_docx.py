#!/usr/bin/env python3
"""Regenerate docs/REPORT.docx (requires: pip install python-docx)."""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

PSEUDO_CODE = """SETUP
  GPIO_init()           // LEDs, push buttons (pull-ups), buzzer, ultrasonic, shift registers, LCD
  USART0_init_9600()
  LCD_init()
  SevenSeg_init()
  Buzzer_init()
  Ultrasound_init()
  PCINT_buttons_init()
  Timer1_init_1Hz_tick()
  sei()

  phase ← GREEN
  sec_left ← 10
  running ← TRUE
  armed ← TRUE
  prev_cm ← large value

MAIN LOOP (forever)
  TrafficLight_step()

TrafficLight_step()
  IF pause_click THEN running ← NOT running
  IF reset_click THEN phase ← GREEN, sec_left ← 10, running ← TRUE; refresh outputs

  IF one_second_tick THEN
      IF violation_lcd_timer > 0 THEN
          decrement timer; IF 0 THEN restore phase text on LCD
      ELSE IF running THEN
          decrement sec_left OR next phase (green→amber→red); refresh 7-segment + LCD

  IF phase = RED AND running THEN
      cm ← Ultrasound_read_cm()
      IF cm = 0 THEN skip
      IF armed AND prev_cm ≥ FAR_cm AND cm ≤ NEAR_cm THEN
          buzzer; LCD VIOLATION; USART print distance
          armed ← FALSE
      IF cm ≥ FAR_cm THEN armed ← TRUE
      prev_cm ← cm"""


def add_report():
    doc = Document()
    title = doc.add_heading(
        'Red-light traffic signal with ultrasonic “camera” (NET3001 Final)', 0
    )
    title.runs[0].font.size = Pt(16)

    doc.add_heading("Title", level=1)
    doc.add_paragraph(
        "Red-light traffic signal with ultrasonic zone detection and USART violation logging."
    )

    doc.add_heading("Name of the group members and contributions", level=1)
    doc.add_paragraph(
        "Abdul Rehman (Student 1): TinkerCAD schematic, demo video, serial testing, "
        "hardware build / bench wiring.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Isaac Abdile (Student 2): Firmware (PlatformIO), GPIO / Timer1 / interrupts / USART, "
        "LCD and 7-segment (shift registers), written report.",
        style="List Bullet",
    )

    doc.add_heading("Brief description of your system", level=1)
    doc.add_paragraph(
        "The traffic light runs 10 s green, 5 s amber, 10 s red, repeating. A 16×2 LCD shows "
        "phase text; a 7-segment display shows seconds remaining via shift registers. Push buttons "
        "pause/resume and reset to green at 10 s. On red, the ultrasonic sensor detects a far-to-near crossing; "
        "the buzzer alerts, the LCD shows VIOLATION, and USART prints VIOLATION dist=XX cm at "
        "9600 baud. Zone limits are constants tuned on the breadboard."
    )

    doc.add_heading("Circuit diagram", level=1)
    p = doc.add_paragraph(
        "[Insert TinkerCAD export (PNG or PDF) here — replace this paragraph.]"
    )
    p.runs[0].italic = True

    doc.add_heading("List of components and topics used", level=1)
    doc.add_paragraph(
        "Components: Arduino Uno; LEDs with resistors; 2 push buttons; 16×2 LCD with contrast "
        "potentiometer; 7-segment + shift registers + resistors; ultrasonic sensor; buzzer; breadboard; jumper wires; USB cable."
    )
    doc.add_paragraph(
        "Topics: GPIO; Timer1 (CTC) for ~1 Hz tick; interrupts (Timer1 compare, pin-change); "
        "USART; ultrasonic sensor ranging; shift registers; LCD 4-bit mode."
    )

    doc.add_heading("Detailed explanation (pseudo-code)", level=1)
    p = doc.add_paragraph(PSEUDO_CODE)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Pt(12)
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    doc.add_heading("Short reflection", level=1)
    doc.add_paragraph(
        "Hardest part: tuning ultrasonic far/near thresholds on real hardware. ISRs stay minimal "
        "(Timer1 sets a flag; work runs in the main loop). Next time: serial calibration mode or "
        "Timer1 input capture on the echo pin."
    )

    out = DOCS / "REPORT.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    DOCS.mkdir(parents=True, exist_ok=True)
    add_report()
